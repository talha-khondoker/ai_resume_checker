"""
File handling and resume extraction utilities.
"""
import os
import mimetypes
from pathlib import Path
from typing import Tuple
import pdfplumber
import PyPDF2
from docx import Document
from app.core.config import settings


def validate_file(
    filename: str,
    file_size: int
) -> Tuple[bool, str]:
    """
    Validate uploaded file.
    
    Args:
        filename: Original filename
        file_size: File size in bytes
    
    Returns:
        Tuple of (is_valid, error_message)
    """
    # Check file extension
    file_ext = Path(filename).suffix.lstrip('.').lower()
    if file_ext not in settings.allowed_extensions_list:
        return False, f"File type .{file_ext} not allowed. Allowed: {', '.join(settings.allowed_extensions_list)}"
    
    # Check file size
    if file_size > settings.MAX_UPLOAD_SIZE:
        max_mb = settings.MAX_UPLOAD_SIZE / (1024 * 1024)
        return False, f"File size exceeds maximum allowed size of {max_mb}MB"
    
    return True, ""


def save_uploaded_file(upload_folder: str, filename: str, content: bytes) -> str:
    """
    Save uploaded file securely.
    
    Args:
        upload_folder: Folder to save file
        filename: Original filename
        content: File content bytes
    
    Returns:
        Relative file path
    """
    # Create upload folder if not exists
    Path(upload_folder).mkdir(parents=True, exist_ok=True)
    
    # Generate safe filename
    import uuid
    file_ext = Path(filename).suffix
    safe_filename = f"{uuid.uuid4()}{file_ext}"
    
    # Full file path
    file_path = os.path.join(upload_folder, safe_filename)
    
    # Save file
    with open(file_path, 'wb') as f:
        f.write(content)
    
    return file_path


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file.
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        Extracted text
    """
    text = ""
    try:
        # Try using pdfplumber first (better for structured documents)
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    except Exception:
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    return text


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        Extracted text
    """
    try:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        return text
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_resume(file_path: str) -> str:
    """
    Extract text from resume file (PDF or DOCX).
    
    Args:
        file_path: Path to resume file
    
    Returns:
        Extracted text
    """
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw extracted text
    
    Returns:
        Cleaned text
    """
    # Remove extra whitespace
    lines = text.split('\n')
    cleaned_lines = [line.strip() for line in lines if line.strip()]
    text = ' '.join(cleaned_lines)
    
    # Remove multiple spaces
    import re
    text = re.sub(r'\s+', ' ', text)
    
    return text
